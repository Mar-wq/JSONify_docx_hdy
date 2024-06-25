import os
# 测试按照章节进行划分的效果
import docx
from docx import Document
from docx.oxml.ns import qn

from jsonfy_docx_xml import jsonfy_book
from jsonfy_docx_xml.elements.paragraph import paragraph
from jsonfy_docx_xml.elements.table import table
from jsonfy_docx_xml.iterators.generic import build_iterators
from jsonfy_docx_xml.textbook.book import Book

#
#


tags_mapping = {qn("w:p"):paragraph
                , qn("w:tbl"):table}


def jsonfy(xml_fragments, document):
    """
    处理单个元素或列表的 XML 片段，并返回相应的 JSON 结构。

    :param xml_fragments: 单个元素或包含元素的列表
    :return: JSON 结构
    """
    if isinstance(xml_fragments, list):
        return [process_single_element(element, document) for element in xml_fragments]
    else:
        return process_single_element(xml_fragments, document)

def process_single_element(element, document):
    """
    处理单个 XML 元素，并返回相应的 JSON 结构。

    :param element: 单个 XML 元素
    :return: JSON 结构
    """
    if element.tag  in tags_mapping:
        return tags_mapping[element.tag](element).to_json(document)

    return {"test": "成功了没"}



def split_document_into_chapters(document):

    body_elements = document.element.body
    chapters = []
    current_chapter = {}

    for element in body_elements:
        if element.tag.endswith('p'):
            if element.style == '1':
                if current_chapter:
                    if current_chapter['sections']:
                        current_chapter['sections'] = split_chapter_into_sections(current_chapter['sections'], document)
                    chapters.append(current_chapter)
                    current_chapter = {}
                current_chapter['chapter_title'] = element.text
                current_chapter['content'] = jsonfy(element, document)
                current_chapter['sections'] = []
                continue

        current_chapter['sections'].append(element)

    if current_chapter:
        if current_chapter['sections']:
            current_chapter['sections'] = split_chapter_into_sections(current_chapter['sections'], document)
        chapters.append(current_chapter)

    return chapters
#
#
# todo  类似于上面的解析章的内容一样，继续解析节的内容
def split_chapter_into_sections(chapter_elements, document):
    sections = []
    current_section = {}

    for element in chapter_elements:
        if element.tag.endswith('p'):
            if element.style == '2':
                if current_section:
                    if current_section['paragraphs']:
                        current_section['paragraphs'] = jsonfy(current_section['paragraphs'], document)
                    sections.append(current_section)
                    current_section = {}
                current_section['section_title'] = element.text
                current_section['content'] = jsonfy(element, document)
                current_section['paragraphs'] = []
                continue

        current_section['paragraphs'].append(element)

    if current_section:
        if current_section['paragraphs']:
            current_section['paragraphs'] = jsonfy(current_section['paragraphs'], document)
        sections.append(current_section)

    return sections
#
#
#
# input_path = r'./docs/第1章 数值分析引论20240226.docx'
# book_title = os.path.splitext(os.path.basename(input_path))[0]
# my_doc = docx.Document(input_path)
# doc_element = my_doc.element
# current =  doc_element.getchildren()[0]
# # 获取 body 元素
# body_element = doc_element.find('.//w:body', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
# all_content =  current.getchildren()
#
# chapters = split_document_into_chapters(input_path)
# document = []
# for chapter in chapters:
#     document.append(split_chapter_into_sections(chapter))
# print()


# build_iterators()
# input_path = r'./docs/第1章 数值分析引论20240226.docx'
# book_title = os.path.splitext(os.path.basename(input_path))[0]
# book_json = {"book_title": book_title}
# document = Document(input_path)
# styles =  document.styles
# # 打印所有样式的名称和ID
# styles_list = []
# for style in styles:
#     styles_list.append({"Style ID": style.style_id, "Style Name": style.name})
# #para0 = document.paragraphs[0]
# book_json['chapters'] = split_document_into_chapters(document)
# print()

build_iterators()
bookPath = r'./docs/第1章 数值分析引论20240226.docx'
document = Document(bookPath)
# book = Book(document)
# book_json = book.to_json()
book_json = jsonfy_book(document)
print()